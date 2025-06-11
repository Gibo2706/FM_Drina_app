from docxtpl import DocxTemplate
from docx import Document
import subprocess
from pathlib import Path
from datetime import datetime
import re
import json



BASE_DIR = Path(__file__).resolve().parent.parent.parent
TEMPLATE_PATH = BASE_DIR / "app" / "templates" / "docx_templates" / "ponuda_template.docx"
OUTPUT_DIR = BASE_DIR / "generated_documents"
OUTPUT_DIR.mkdir(exist_ok=True)

def generisi_ponudu(data: dict) -> str:
    doc = DocxTemplate(TEMPLATE_PATH)
    doc.render(data)

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    safe_name = data['klijent_naziv'].replace(" ", "_").replace(".", "").replace(",", "")
    naziv_fajla = f"Ponuda_{safe_name}_{timestamp}.docx"
    putanja = OUTPUT_DIR / naziv_fajla

    doc.save(putanja)

    meta = {
        "tip": "ponuda",
        "broj": data.get("broj_ponude"),
        "klijent": data.get("klijent_naziv"),
        "datum": data.get("datum"),
        "fajl": naziv_fajla
    }

    with open(putanja.with_suffix('.json'), "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    return naziv_fajla  # samo ime fajla, ne puna putanja


def generisi_broj_ponude():
    godina = datetime.now().year % 100  # npr. 25 za 2025.
    folder = Path(__file__).resolve().parent.parent.parent / "generated_documents"

    broj = 1
    pattern = re.compile(r'Ponuda_.*_(\d{8})\d{6}\.docx')

    for f in folder.glob("Ponuda_*.docx"):
        match = pattern.match(f.name)
        if match:
            broj += 1

    return f"{broj}/{godina}"

UGOVOR_TEMPLATE_PATH = BASE_DIR / "app" / "templates" / "docx_templates" / "ugovor_template_full.docx"

def generisi_ugovor(data: dict) -> str:
    doc = Document(UGOVOR_TEMPLATE_PATH)

    # Zamena placeholder-a u paragrafima
    for p in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{ {key} }}}}" in p.text:
                p.text = p.text.replace(f"{{{{ {key} }}}}", str(value))

    # Zamena u tabelama ako ih ima
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if f"{{{{ {key} }}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{ {key} }}}}", str(value))

    # Naziv fajla
    klijent = data.get("narucilac_naziv", "Ugovor").replace(" ", "_").upper()
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    naziv_fajla = f"Ugovor_{klijent}_{timestamp}.docx"
    putanja = OUTPUT_DIR / naziv_fajla

    doc.save(putanja)

    # Snimi i .json metapodatke
    meta = {
        "tip": "ugovor",
        "klijent": data.get("narucilac_naziv"),
        "datum": data.get("datum"),
        "fajl": naziv_fajla,
        "povezana_ponuda": data.get("povezana_ponuda", "")
    }

    with open(putanja.with_suffix('.json'), "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)

    # Pokušaj konverzije u PDF
    try:
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf",
            "--outdir", str(putanja.parent), str(putanja)
        ], check=True)
    except Exception as e:
        print(f"[PDF konverzija nije uspela] {e}")

    return naziv_fajla